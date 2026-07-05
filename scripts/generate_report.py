#!/usr/bin/env python
"""Génère le dossier de projet (rapport Word .docx), niveau titre RNCP.

Produit un document professionnel et structuré (page de garde, sommaire
automatique, pied de page numéroté, sections détaillées couvrant les
compétences C9 à C13, difficultés rencontrées, conclusion, annexes).

Réutilise les captures de docs/fr/captures/, les figures de reports/figures/
et les métriques de reports/metrics.json.

Lancer :  ./.venv/bin/python scripts/generate_report.py
Sortie :  docs/fr/RAPPORT_C9-C13.docx
"""
from __future__ import annotations

import json
import struct
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CAP = ROOT / "docs" / "fr" / "captures"
FIG = ROOT / "reports" / "figures"
OUT = ROOT / "docs" / "fr" / "RAPPORT_C9-C13.docx"

MAX_W_IN = 5.0
MAX_H_IN = 3.4
ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # bleu nuit pour les titres
GREY = RGBColor(0x55, 0x55, 0x55)

metrics = json.loads((ROOT / "reports" / "metrics.json").read_text())
BEST = metrics["best_model"]
RESULTS = metrics["results"]
N_FEATURES = len(metrics["features"])


# --------------------------------------------------------------------------
# Helpers bas niveau
# --------------------------------------------------------------------------
def png_size(path: Path) -> tuple[int, int]:
    with open(path, "rb") as fh:
        head = fh.read(24)
    if head[:8] != b"\x89PNG\r\n\x1a\n":
        return (1, 1)
    return struct.unpack(">II", head[16:24])


def body(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_image(doc, path: Path, caption: str):
    if not path.exists():
        body(doc, f"[capture manquante : {path.name}]")
        return
    w, h = png_size(path)
    if h / w * MAX_W_IN > MAX_H_IN:
        doc.add_picture(str(path), height=Inches(MAX_H_IN))
    else:
        doc.add_picture(str(path), width=Inches(MAX_W_IN))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.runs[0]
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def add_mono(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def quote(doc, text):
    p = doc.add_paragraph(text)
    try:
        p.style = doc.styles["Intense Quote"]
    except KeyError:
        for r in p.runs:
            r.italic = True


def add_field(run, instr, placeholder=""):
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = placeholder
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f2, t, f3):
        run._r.append(el)


def add_toc(doc):
    p = doc.add_paragraph()
    add_field(p.add_run(), 'TOC \\o "1-3" \\h \\z \\u',
              "[ Sommaire : clic droit › Mettre à jour les champs (ou Ctrl+A puis F9) ]")


def set_update_fields(doc):
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Dossier de projet : Compétences C9 à C13      Page ").font.size = Pt(8)
    add_field(p.add_run(), "PAGE")
    p.add_run(" / ").font.size = Pt(8)
    add_field(p.add_run(), "NUMPAGES")


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


# --------------------------------------------------------------------------
# Construction du document
# --------------------------------------------------------------------------
def cover(doc):
    """Page de garde reprise du Rapport_E4 (SmartTicket), avec notre projet."""
    def line(text, size, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        return p

    def blank(n=1):
        for _ in range(n):
            doc.add_paragraph()

    blank(3)
    line("RAPPORT DE PROJET", 28, bold=True)
    line("Prédiction de la qualité de l’air", 20, bold=True)
    line("Solution d’intelligence artificielle : de la donnée au service "
         "en production (compétences C9 à C13)", 12)
    blank(4)
    line("Présenté par", 10)
    line("MARIMOUN MOHAMED", 16, bold=True)
    blank(1)
    line("Titre RNCP 37827", 10)
    line("Développeur en Intelligence Artificielle, Niveau 6 (Bac+3)", 10, bold=True)
    line("ÉPREUVE E4", 10, bold=True)
    blank(1)
    line("Année 2025-2026", 10)
    doc.add_page_break()


def sommaire(doc):
    h1(doc, "Sommaire")
    add_toc(doc)
    doc.add_page_break()


def introduction(doc):
    h1(doc, "1. Introduction")

    h2(doc, "1.1 Contexte du projet")
    body(doc,
         "La qualité de l’air est un enjeu de santé publique et d’aménagement "
         "urbain majeur. La capacité à estimer en continu un indice de qualité de "
         "l’air (Air Quality Index, AQI) à partir de capteurs permet d’émettre des "
         "alertes, d’informer les décisions publiques et de réduire l’exposition "
         "des populations. Ce projet ne se limite pas à un modèle "
         "prédictif : il couvre toute la chaîne logicielle qui permet de "
         "l’exploiter en production de façon fiable et maintenable.")
    body(doc,
         "Le jeu de données de référence (AirQualityData.csv) contient 4 000 "
         "relevés horaires de concentrations de polluants (CO, NOx, NO2, O3, SO2, "
         "PM2.5, PM10) et de variables météorologiques (température, humidité, "
         "pression, vent), ainsi que des variables dérivées et temporelles.")

    h2(doc, "1.2 Problématique")
    body(doc,
         "Au-delà de la modélisation, le défi industriel consiste à rendre le "
         "modèle exploitable par les autres composants d’un système d’information : "
         "comment exposer le modèle de manière standardisée et sécurisée ? Comment "
         "l’intégrer dans une application accessible à un utilisateur métier ? "
         "Comment surveiller son comportement dans le temps et détecter une "
         "dégradation ? Comment garantir un niveau de qualité élevé par des tests "
         "automatisés ? Et comment automatiser la validation, le test, le "
         "packaging et le déploiement dans une démarche MLOps ? Ces questions "
         "correspondent précisément aux compétences C9 à C13 traitées dans ce "
         "dossier.")

    h2(doc, "1.3 Objectifs")
    bullet(doc, "Développer une API REST exposant le modèle, conforme à ses "
                "spécifications fonctionnelles et techniques et aux standards de "
                "qualité et de sécurité (C9).")
    bullet(doc, "Intégrer cette API dans une application respectant les normes "
                "d’accessibilité (C10).")
    bullet(doc, "Mettre en place le monitoring du modèle : collecte de métriques, "
                "alerte et restitution (C11).")
    bullet(doc, "Programmer les tests automatisés couvrant données, préparation, "
                "entraînement, évaluation et validation (C12).")
    bullet(doc, "Créer une chaîne de livraison continue dans une approche MLOps "
                "(C13).")

    h2(doc, "1.4 Périmètre et plan du dossier")
    body(doc,
         "Le présent dossier présente d’abord les données et la démarche "
         "scientifique (section 2), puis l’architecture technique de la solution "
         "(section 3). Il détaille ensuite la réalisation de chacune des cinq "
         "compétences C9 à C13 (sections 4 à 8), en s’appuyant sur des éléments de "
         "preuve (captures d’écran et sorties de programme). Les difficultés "
         "rencontrées et leurs solutions sont analysées en section 9, avant la "
         "conclusion (section 10) et les annexes.")


def donnees(doc):
    h1(doc, "2. Données et démarche scientifique")

    h2(doc, "2.1 Le jeu de données")
    body(doc,
         "Le fichier AirQualityData.csv comporte 4 000 observations horaires et 23 "
         "colonnes : sept polluants, cinq variables météorologiques, des variables "
         "dérivées (ratios entre polluants, moyennes mobiles sur trois pas, indice "
         "température-humidité), des variables temporelles (heure, jour de la "
         "semaine) et une colonne cible fournie, AirQualityIndex.")

    h2(doc, "2.2 Une découverte déterminante : une cible non apprenable")
    body(doc,
         "En explorant les données, j’ai constaté un problème : la colonne cible "
         "AirQualityIndex est indépendante de toutes les variables explicatives "
         "(corrélation de Pearson inférieure à 0,04 en valeur absolue pour chaque "
         "colonne). Cette cible se comporte donc comme un bruit aléatoire : aucun "
         "modèle ne peut l’apprendre. Les trois familles de modèles testées "
         "(régression linéaire, forêt aléatoire, gradient boosting) obtiennent "
         "toutes un R² négatif, c’est-à-dire une performance inférieure à la simple "
         "prédiction de la moyenne.")

    h2(doc, "2.3 La solution retenue : recalculer un AQI normalisé (EPA)")
    body(doc,
         "Plutôt que d’optimiser en vain un modèle sur une cible sans signal, j’ai "
         "reconstruit une cible qui a du sens : l’indice de qualité de l’air "
         "officiel de l’Agence américaine de protection de l’environnement (EPA). "
         "Cet indice est une fonction déterministe, linéaire par morceaux, des "
         "concentrations de polluants. Pour chaque polluant, la concentration est "
         "convertie en sous-indice à l’aide d’une table de seuils, et l’AQI global "
         "est le maximum des sous-indices (le « polluant dominant » pilote "
         "l’indice). J’ai implémenté ce calcul dans src/air_quality/aqi.py, et "
         "j’ai écarté la colonne bruitée d’origine pour qu’elle ne contamine jamais "
         "les variables explicatives.")

    h2(doc, "2.4 Résultats de modélisation")
    body(doc,
         f"Sur la cible AQI recalculée, les modèles retrouvent une structure forte. "
         f"Le meilleur modèle est sélectionné automatiquement sur le R² de test ; "
         f"il s’agit ici de « {BEST} ». Le tableau suivant compare les trois "
         f"modèles entraînés sur {N_FEATURES} variables.")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, txt in enumerate(["Modèle", "R² (validation croisée)", "R² (test)", "RMSE (test)"]):
        table.rows[0].cells[i].paragraphs[0].add_run(txt).bold = True
    lib = {"linear_regression": "Régression linéaire",
           "random_forest": "Forêt aléatoire (Random Forest)", "xgboost": "XGBoost"}
    for name, m in RESULTS.items():
        cells = table.add_row().cells
        nom = lib.get(name, name) + ("  (meilleur)" if name == BEST else "")
        cells[0].text = nom
        cells[1].text = f"{m['cv_r2_mean']:.4f}"
        cells[2].text = f"{m['test_r2']:.4f}"
        cells[3].text = f"{m['test_rmse']:.3f}"
    doc.add_paragraph()
    body(doc,
         "La forêt aléatoire atteint un R² de test de l’ordre de 0,999, avec une "
         "erreur quadratique moyenne très faible au regard de l’échelle de l’AQI "
         "(0 à 500). Point important : les variables les plus influentes sont "
         "l’ozone (O3), les PM2.5 et les PM10, exactement les polluants qui "
         "pilotent la formule EPA. Le modèle a donc appris une logique physique "
         "réelle.")
    add_image(doc, FIG / "pred_vs_actual.png",
              "Figure 1 : Valeurs prédites vs valeurs réelles sur le jeu de test : "
              "les points s’alignent sur la diagonale.")
    add_image(doc, FIG / "feature_importance.png",
              "Figure 2 : Importance des variables : O3, PM2.5 et PM10 dominent, "
              "ce qui valide la cohérence du modèle avec la définition de l’AQI.")

    h2(doc, "2.5 Garanties méthodologiques")
    bullet(doc, "Absence de fuite de données : la normalisation est encapsulée "
                "dans un Pipeline scikit-learn et ajustée uniquement sur le pli "
                "d’entraînement ; l’ingénierie des variables est sans état.")
    bullet(doc, "Déterminisme : une graine aléatoire unique (RANDOM_STATE) "
                "garantit que l’entraînement et l’évaluation partagent exactement "
                "le même découpage entre exécutions.")
    bullet(doc, "Reproductibilité : les hyperparamètres sont externalisés dans "
                "params.yaml et le pipeline est orchestré par DVC.")

    h2(doc, "2.6 Calcul détaillé de l’AQI")
    body(doc,
         "Pour chaque polluant, la concentration mesurée C est située dans un "
         "intervalle de seuils [BP_lo, BP_hi] associé à un intervalle d’indice "
         "[I_lo, I_hi]. Le sous-indice est obtenu par interpolation linéaire, puis "
         "l’AQI global est le maximum des sous-indices :")
    add_mono(doc,
             "I_p = (I_hi - I_lo) / (BP_hi - BP_lo) * (C - BP_lo) + I_lo\n"
             "AQI = max(I_p)   sur l’ensemble des polluants")
    body(doc,
         "Le polluant qui atteint ce maximum est le « polluant dominant ». "
         "Exemple de table de seuils EPA pour les PM2.5 (µg/m³) :")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    for i, t in enumerate(["Concentration PM2.5 (µg/m³)", "Indice AQI", "Catégorie"]):
        tbl.rows[0].cells[i].paragraphs[0].add_run(t).bold = True
    for c, idx, cat in [("0,0 à 12,0", "0 à 50", "Bon"),
                        ("12,1 à 35,4", "51 à 100", "Modéré"),
                        ("35,5 à 55,4", "101 à 150", "Mauvais pour groupes sensibles"),
                        ("55,5 à 150,4", "151 à 200", "Mauvais")]:
        cells = tbl.add_row().cells
        cells[0].text = c
        cells[1].text = idx
        cells[2].text = cat
    doc.add_paragraph()
    body(doc,
         "Hypothèses d’unités pour ce jeu de données : PM2.5 et PM10 en µg/m³ "
         "(directement) ; CO en ppm ; O3, NO2 et SO2 en ppb (O3 converti en ppm "
         "pour la table 8 heures). Le NOx, qui n’est pas un polluant AQI officiel, "
         "est conservé comme variable explicative mais exclu du calcul de la cible.")

    h2(doc, "2.7 Ingénierie des variables")
    body(doc,
         "Les variables sont enrichies de façon sans état : les mêmes "
         "transformations s’appliquent à l’entraînement et à l’inférence. Trois "
         "familles sont ajoutées :")
    bullet(doc, "Calendaires : mois et saison extraits de l’horodatage.")
    bullet(doc, "Encodages cycliques : l’heure, le mois et la direction du vent "
                "sont projetés en sinus et cosinus pour respecter la continuité "
                "(23 h proche de 0 h, décembre proche de janvier).")
    bullet(doc, "Interactions physiques : PM2.5 × Température et O3 × Humidité.")
    add_mono(doc,
             "Hour_sin = sin(2*pi*Hour/24)      Hour_cos = cos(2*pi*Hour/24)\n"
             "Wind_sin = sin(WindDirection)     Wind_cos = cos(WindDirection)")

    h2(doc, "2.8 Méthodologie de modélisation")
    body(doc,
         "Le jeu est découpé en 80 % d’entraînement et 20 % de test. Chaque modèle "
         "est évalué par validation croisée à 5 plis sur l’entraînement, puis le "
         "meilleur est retenu sur le R² de test. Chaque modèle est encapsulé dans "
         "un Pipeline scikit-learn afin que la normalisation soit ajustée uniquement "
         "sur le pli d’entraînement (protection contre la fuite de données) :")
    add_mono(doc,
             "Pipeline([\n"
             "    ('scaler', StandardScaler()),\n"
             "    ('model', estimateur),\n"
             "])")
    body(doc, "Hyperparamètres principaux (externalisés dans params.yaml) :")
    import yaml
    params = yaml.safe_load((ROOT / "params.yaml").read_text())
    rf = params["models"]["random_forest"]
    xgb = params["models"]["xgboost"]
    sp = params["split"]
    hp = [
        ("Taille du jeu de test", f"{int(sp['test_size'] * 100)} %"),
        ("Validation croisée", f"{sp['cv_folds']} plis"),
        ("Graine aléatoire (RANDOM_STATE)", str(sp["random_state"])),
        ("Random Forest, n_estimators", str(rf["n_estimators"])),
        ("XGBoost, n_estimators", str(xgb["n_estimators"])),
        ("XGBoost, learning_rate", str(xgb["learning_rate"])),
        ("XGBoost, max_depth", str(xgb["max_depth"])),
        ("XGBoost, subsample", str(xgb["subsample"])),
        ("XGBoost, colsample_bytree", str(xgb["colsample_bytree"])),
    ]
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Light Grid Accent 1"
    t2.rows[0].cells[0].paragraphs[0].add_run("Paramètre").bold = True
    t2.rows[0].cells[1].paragraphs[0].add_run("Valeur").bold = True
    for k, v in hp:
        cells = t2.add_row().cells
        cells[0].text = k
        cells[1].text = v
    doc.add_paragraph()
    doc.add_page_break()


def architecture(doc):
    h1(doc, "3. Architecture technique de la solution")
    body(doc,
         "La solution suit une architecture modulaire en couches, organisée autour "
         "d’un paquet Python installable (src/air_quality/). Le principe directeur "
         "est la non-duplication de la logique métier : l’API et l’application "
         "partagent un unique chemin d’inférence (module inference.py), ce qui "
         "interdit toute divergence de la préparation des variables entre les deux "
         "interfaces.")

    h2(doc, "3.1 Couche « laboratoire » (préparation et modélisation)")
    bullet(doc, "aqi.py : calcul de l’AQI EPA (tables de seuils, sous-indices).")
    bullet(doc, "data.py : chargement, validation et constitution du jeu de données.")
    bullet(doc, "features.py : ingénierie des variables (calendaires, cycliques, "
                "interactions).")
    bullet(doc, "modeling.py : entraînement, validation croisée, sélection et "
                "sauvegarde du meilleur modèle.")
    bullet(doc, "evaluate.py : métriques et graphiques d’évaluation.")

    h2(doc, "3.2 Couche de service et d’exploitation")
    bullet(doc, "inference.py : featurisation partagée et fonction de prédiction "
                "unitaire.")
    bullet(doc, "api/ : service REST FastAPI (main, schemas, security).")
    bullet(doc, "app/streamlit_app.py : application cliente.")
    bullet(doc, "validation.py : règles de validation des données (pandera).")
    bullet(doc, "monitoring.py et logging_config.py : supervision, dérive, "
                "journalisation.")

    h2(doc, "3.3 Flux de données")
    body(doc,
         "À l’entraînement : le CSV est chargé et validé, l’AQI EPA est calculé, "
         "les variables sont enrichies, les données sont découpées puis chaque "
         "modèle est entraîné dans un pipeline incluant la normalisation ; le "
         "meilleur modèle et ses métriques sont persistés. À l’inférence : une "
         "requête (mesures + horodatage) est validée par les schémas Pydantic, "
         "transformée en vecteur de variables par inference.py, qui réutilise "
         "exactement la même ingénierie que l’entraînement, puis soumise au "
         "modèle ; la réponse comprend l’AQI, sa catégorie et le polluant dominant.")
    doc.add_page_break()


# --------------------------------------------------------------------------
# Sections compétences
# --------------------------------------------------------------------------
PYTEST = ("$ pytest -m 'not slow'      ->  32 passed\n"
          "$ pytest -m slow             ->   1 passed   (barriere qualite : R2 > 0.9)\n"
          "Total : 33 tests verts, rejoues automatiquement par la CI.")


def comp_c9(doc):
    h1(doc, "4. C9 : Développer une API REST exposant le modèle")
    body(doc, "Compétence visée :")
    quote(doc,
          "Développer une API REST exposant un modèle d’intelligence artificielle "
          "en respectant ses spécifications fonctionnelles et techniques et les "
          "standards de qualité et de sécurité du marché pour permettre "
          "l’interaction entre le modèle et les autres composants du projet.")

    h2(doc, "4.1 Choix technologique")
    body(doc,
         "J’ai développé le service avec FastAPI, servi par uvicorn. J’ai retenu "
         "ce framework parce qu’il valide nativement les entrées (Pydantic), génère "
         "automatiquement la documentation OpenAPI, et fournit les briques de "
         "sécurité dont j’avais besoin (CORS, limitation de débit).")

    h2(doc, "4.2 Spécifications fonctionnelles")
    body(doc, "L’API expose sept routes répondant aux besoins d’interaction avec "
              "les autres composants :")
    bullet(doc, "GET /health : sonde de vivacité et indicateur de modèle chargé.")
    bullet(doc, "GET /model/info : métadonnées du modèle servi et métriques "
                "d’entraînement.")
    bullet(doc, "POST /predict : prédiction unitaire à partir d’une observation.")
    bullet(doc, "POST /predict/batch : prédiction par lot (jusqu’à 1 000 éléments).")
    bullet(doc, "GET /metrics : exposition Prometheus pour la supervision.")
    bullet(doc, "GET /docs et /redoc : documentation interactive de l’API.")

    h2(doc, "4.3 Spécifications techniques et contrat d’interface")
    body(doc,
         "Les schémas Pydantic (api/schemas.py) constituent le contrat d’entrée et "
         "de sortie. Chaque champ de mesure est typé et borné par des contraintes "
         "de plage (par exemple l’humidité entre 0 et 100 %). Le chargement du "
         "modèle s’effectue une seule fois au démarrage du service (cycle de vie "
         "« lifespan »), afin de garantir des temps de réponse faibles. La "
         "documentation technique est générée automatiquement et servie sur /docs ; "
         "elle est complétée par une spécification rédigée, docs/api_spec.md.")
    body(doc, "Exemple d’appel et de réponse de la route /predict :")
    add_mono(doc,
             "POST /predict      (en-tete X-API-Key requis)\n"
             "{\n"
             '  "co": 3.8, "no2": 144.3, "o3": 118.1, "pm25": 147.3,\n'
             '  "pm10": 208.8, "temperature": 28.5, "humidity": 45\n'
             "}\n"
             "Reponse 200 :\n"
             '{ "aqi": 213.98, "category": "Tres mauvais",\n'
             '  "dominant_pollutant": "O3(GT)" }')
    body(doc, "Codes de réponse renvoyés par l’API :")
    codes = [("200", "Succès"),
             ("401", "Clé d’API manquante ou invalide"),
             ("422", "Entrée invalide ou hors plage"),
             ("429", "Limite de débit dépassée"),
             ("500", "Erreur interne")]
    tc = doc.add_table(rows=1, cols=2)
    tc.style = "Light Grid Accent 1"
    tc.rows[0].cells[0].paragraphs[0].add_run("Code HTTP").bold = True
    tc.rows[0].cells[1].paragraphs[0].add_run("Signification").bold = True
    for code, sens in codes:
        cells = tc.add_row().cells
        cells[0].text = code
        cells[1].text = sens
    doc.add_paragraph()

    h2(doc, "4.4 Standards de qualité et de sécurité")
    bullet(doc, "Validation des entrées : toute requête malformée ou hors plage "
                "est rejetée avec un code HTTP 422 avant d’atteindre le modèle.")
    bullet(doc, "Authentification : clé d’API transmise dans l’en-tête X-API-Key ; "
                "une clé absente ou invalide entraîne un code HTTP 401.")
    bullet(doc, "Contrôle d’origine (CORS) configurable et limitation de débit "
                "(60 requêtes/minute sur /predict) pour prévenir les abus.")
    bullet(doc, "Gestion centralisée des erreurs et journalisation structurée.")

    h2(doc, "4.5 Éléments de preuve")
    add_image(doc, CAP / "c9_swagger_overview.png",
              "Figure 3 : Documentation OpenAPI (Swagger) générée automatiquement.")
    add_image(doc, CAP / "c9_model_info.png",
              "Figure 4 : Réponse JSON de GET /model/info (métadonnées du modèle).")

    h2(doc, "4.6 Bilan")
    body(doc,
         "L’API est l’interface entre le modèle et le reste du système. Elle "
         "respecte ses spécifications fonctionnelles et techniques et applique les "
         "standards de qualité et de sécurité courants du marché.")
    doc.add_page_break()


def comp_c10(doc):
    h1(doc, "5. C10 : Intégrer l’API dans une application")
    body(doc, "Compétence visée :")
    quote(doc,
          "Intégrer l’API d’un modèle ou d’un service d’intelligence artificielle "
          "dans une application, en respectant les spécifications du projet et les "
          "normes d’accessibilité en vigueur, à l’aide de la documentation "
          "technique de l’API, afin de créer les fonctionnalités d’intelligence "
          "artificielle de l’application.")

    h2(doc, "5.1 Application réalisée")
    body(doc,
         "J’ai développé une application web avec Streamlit (app/streamlit_app.py). "
         "Elle appelle l’API avec la bibliothèque httpx, en suivant le contrat "
         "OpenAPI de la route /predict. Elle offre deux fonctionnalités : la "
         "prédiction unitaire à partir d’un formulaire, et la prédiction par lot "
         "par import d’un fichier CSV. L’adresse de l’API et la clé sont fournies "
         "par des variables d’environnement, hors du code.")

    h2(doc, "5.2 Respect des normes d’accessibilité")
    body(doc,
         "J’ai traité l’accessibilité (référentiels RGAA et WCAG) de façon "
         "explicite. La catégorie de qualité de l’air est restituée par trois "
         "canaux (texte, icône et couleur), jamais par la couleur seule, "
         "conformément au critère WCAG 1.4.1 sur l’usage de la couleur. "
         "Chaque champ de saisie possède un libellé explicite et une aide "
         "contextuelle, et la page respecte une hiérarchie de titres.")
    body(doc,
         "Une limite est assumée et documentée dans l’application : Streamlit offre "
         "un contrôle restreint sur certains aspects fins du RGAA (ordre de "
         "tabulation, attributs ARIA). Le cas échéant, une page HTML dédiée, "
         "strictement conforme, pourrait réutiliser la même route /predict sans "
         "modification du backend, preuve de la bonne séparation des "
         "responsabilités.")

    h2(doc, "5.3 Éléments de preuve")
    add_image(doc, CAP / "c10_form.png",
              "Figure 5 : Formulaire de saisie de l’application (libellés et aides).")
    add_image(doc, CAP / "c10_result.png",
              "Figure 6 : Restitution du résultat : AQI, catégorie (texte + icône + "
              "couleur) et polluant dominant.")
    doc.add_page_break()


def comp_c11(doc):
    h1(doc, "6. C11 : Monitorer le modèle")
    body(doc, "Compétence visée :")
    quote(doc,
          "Monitorer un modèle d’intelligence artificielle à partir des métriques "
          "courantes et spécifiques au projet, en intégrant les outils de collecte, "
          "d’alerte et de restitution des données du monitorage pour permettre "
          "l’amélioration du modèle de façon itérative.")

    h2(doc, "6.1 Collecte des métriques")
    body(doc,
         "L’API expose un point de terminaison /metrics au format Prometheus, "
         "standard du marché pour la supervision. Aux métriques courantes (volume "
         "de requêtes, latence, taux d’erreur) s’ajoutent des métriques spécifiques "
         "au projet : nombre de prédictions par catégorie de qualité de l’air "
         "(aqi_predictions_total), distribution des valeurs d’AQI prédites et "
         "compteur d’erreurs de prédiction. Par ailleurs, chaque prédiction servie "
         "est journalisée au format JSON dans reports/predictions_log.jsonl, ce qui "
         "permet d’analyser la dérive sur le trafic réel.")

    h2(doc, "6.2 Détection de dérive et restitution")
    body(doc,
         "La dérive des données (data drift) est mesurée par un test statistique de "
         "Kolmogorov-Smirnov appliqué colonne par colonne, comparant les données de "
         "référence (entraînement) à un lot courant. Ce choix garantit un signal "
         "stable, indépendant des évolutions d’outils tiers. Un rapport HTML riche "
         "est par ailleurs produit avec Evidently pour la restitution visuelle.")

    h2(doc, "6.3 Mécanisme d’alerte")
    body(doc,
         "Lorsque la proportion de colonnes en dérive dépasse un seuil défini dans "
         "params.yaml, une alerte est écrite dans reports/alerts/ et la commande se "
         "termine avec un code de sortie non nul, exploitable par une tâche "
         "planifiée ou par la chaîne d’intégration continue. Ce dispositif a été "
         "validé en injectant volontairement une perturbation des capteurs "
         "(commande « monitor --perturb »), ce qui déclenche l’alerte comme attendu.")
    body(doc,
         "Une colonne est considérée en dérive lorsque la p-value du test KS est "
         "inférieure à 0,05. Exemple d’alerte écrite lors du test de perturbation :")
    add_mono(doc,
             "reports/alerts/drift_alert.json :\n"
             '{ "drift_share": 0.57, "threshold": 0.5,\n'
             '  "n_drifted": 12, "n_columns": 21 }')


def comp_c12(doc):
    h1(doc, "7. C12 : Programmer les tests automatisés")
    body(doc, "Compétence visée :")
    quote(doc,
          "Programmer les tests automatisés d’un modèle d’intelligence "
          "artificielle en définissant les règles de validation des jeux de "
          "données, des étapes de préparation des données, d’entraînement, "
          "d’évaluation et de validation du modèle pour permettre son intégration "
          "en continu et garantir un niveau de qualité élevé.")

    h2(doc, "7.1 Stratégie de test")
    body(doc,
         "Une suite de tests automatisés (dossier tests/, framework pytest) couvre "
         "l’ensemble de la chaîne, de la donnée brute jusqu’au modèle validé. Les "
         "tests sont organisés en deux catégories : une suite rapide, exécutée à "
         "chaque modification, et un test lent de « barrière de qualité » qui "
         "entraîne réellement le modèle.")

    h2(doc, "7.2 Règles de validation des jeux de données")
    body(doc,
         "Les règles de validation des données sont formalisées avec la "
         "bibliothèque pandera (module validation.py) : types attendus et plages "
         "physiquement plausibles (concentrations positives, humidité comprise "
         "entre 0 et 100 %, etc.). Ces règles sont vérifiées à deux endroits : à "
         "l’exécution du pipeline (dans data.build_dataset) et dans les tests, ce "
         "qui garantit qu’une donnée corrompue échoue au plus tôt.")
    rules = [("PM2.5, PM10, CO, NOx, NO2, O3, SO2", "≥ 0"),
             ("Humidity", "entre 0 et 100"),
             ("WindDirection", "entre 0 et 360"),
             ("Pressure", "entre 870 et 1085"),
             ("Hour", "entre 0 et 23"),
             ("DayOfWeek", "entre 0 et 6"),
             ("AQI (cible)", "entre 0 et 500")]
    tr = doc.add_table(rows=1, cols=2)
    tr.style = "Light Grid Accent 1"
    tr.rows[0].cells[0].paragraphs[0].add_run("Colonne").bold = True
    tr.rows[0].cells[1].paragraphs[0].add_run("Contrainte").bold = True
    for col, cons in rules:
        cells = tr.add_row().cells
        cells[0].text = col
        cells[1].text = cons
    doc.add_paragraph()

    h2(doc, "7.3 Couverture des étapes")
    bullet(doc, "Préparation : calcul de l’AQI EPA (points de référence connus), "
                "ingénierie des variables (encodages cycliques, colonnes attendues).")
    bullet(doc, "Entraînement : forme des découpages, cohérence des variables.")
    bullet(doc, "Validation du modèle : un test échoue si le meilleur modèle "
                "n’atteint pas un R² supérieur à 0,9 sur le jeu de test.")
    bullet(doc, "Service : tests d’intégration de l’API (codes 200, 401, 422).")
    body(doc, "Organisation des fichiers de test :")
    bullet(doc, "test_aqi.py : sous-indices EPA sur des points de référence connus.")
    bullet(doc, "test_data.py et test_features.py : préparation et ingénierie.")
    bullet(doc, "test_validation.py : règles pandera (cas valides et corrompus).")
    bullet(doc, "test_modeling.py : découpage et barrière de qualité (R² > 0,9).")
    bullet(doc, "test_api.py : intégration de l’API via le client de test FastAPI.")

    h2(doc, "7.4 Élément de preuve")
    add_mono(doc, PYTEST)
    body(doc,
         "L’intégralité de la suite est rejouée automatiquement par la chaîne "
         "d’intégration continue (voir compétence C13), ce qui garantit un niveau "
         "de qualité élevé et constant à chaque évolution du code.")


def comp_c13(doc):
    h1(doc, "8. C13 : Chaîne de livraison continue (MLOps)")
    body(doc, "Compétence visée :")
    quote(doc,
          "Créer une chaîne de livraison continue d’un modèle d’intelligence "
          "artificielle en installant les outils et en appliquant les "
          "configurations souhaitées, dans le respect du cadre imposé par le "
          "projet et dans une approche MLOps, pour automatiser les étapes de "
          "validation, de test, de packaging et de déploiement du modèle.")

    h2(doc, "8.1 Pipeline reproductible (DVC)")
    body(doc,
         "Le cycle de production du modèle est décrit comme un pipeline DVC "
         "(dvc.yaml) paramétré par params.yaml. Le graphe d’étapes "
         "build_target → train → evaluate est reproductible par la commande "
         "« dvc repro », qui ne recalcule que les étapes dont les dépendances ont "
         "changé et qui versionne les données et le modèle produits.")
    add_mono(doc,
             "dvc.yaml (étapes) :\n"
             "  build_target : deps  AirQualityData.csv, aqi.py, data.py\n"
             "                 outs  data/air_quality_with_aqi.csv\n"
             "  train        : params split, models\n"
             "                 outs  models/best_model.joblib, reports/metrics.json\n"
             "  evaluate     : outs  reports/figures/*.png")

    h2(doc, "8.2 Packaging")
    body(doc,
         "Le service est empaqueté dans une image Docker (Dockerfile multi-étapes, "
         "utilisateur non privilégié). Le fichier docker-compose.yml permet de "
         "lancer simultanément l’API et l’application à partir de la même image, ce "
         "qui garantit une exécution identique sur n’importe quel environnement.")

    h2(doc, "8.3 Intégration et déploiement continus")
    body(doc,
         "Deux workflows GitHub Actions automatisent la chaîne. Le workflow "
         "d’intégration continue (ci.yml) enchaîne, à chaque poussée de code : "
         "analyse statique (ruff), entraînement, exécution des tests, puis "
         "construction de l’image. Le workflow de livraison (cd.yml) reproduit le "
         "modèle, publie l’image sur un registre (GHCR) et archive l’artefact "
         "modèle lors d’une mise en tag. Un dispositif pre-commit applique enfin le "
         "linter avant chaque validation locale.")
    body(doc, "Étapes du workflow d’intégration continue (ci.yml) :")
    bullet(doc, "Installer le paquet et ses dépendances.")
    bullet(doc, "Analyse statique du code (ruff).")
    bullet(doc, "Construire la cible AQI puis entraîner le modèle.")
    bullet(doc, "Exécuter les tests, y compris la barrière de qualité.")
    bullet(doc, "Construire l’image Docker.")
    body(doc,
         "L’ensemble matérialise une démarche MLOps : automatisation des étapes de "
         "validation, de test, de packaging et de déploiement, avec traçabilité des "
         "données et du modèle.")


def difficultes(doc):
    h1(doc, "9. Difficultés rencontrées et solutions")

    h2(doc, "9.1 Une cible de modélisation non apprenable")
    body(doc,
         "Ma principale difficulté a été de diagnostiquer une cible fournie sans "
         "signal exploitable. J’ai choisi de reconstruire une cible normalisée et "
         "déterministe (AQI EPA) plutôt que de forcer un modèle ; le problème est "
         "redevenu apprenable sans trahir le sens métier.")

    h2(doc, "9.2 Évolution rapide d’un outil tiers (Evidently)")
    body(doc,
         "La bibliothèque Evidently a beaucoup changé d’interface entre versions. "
         "Pour ne pas rendre l’alerte dépendante d’une API instable, j’ai "
         "ré-implémenté la détection de dérive avec un test statistique que je "
         "maîtrise (Kolmogorov-Smirnov) ; je ne me sers d’Evidently que pour la "
         "restitution HTML, de manière défensive.")

    h2(doc, "9.3 Conflit d’interpréteur Python avec DVC")
    body(doc,
         "L’outil DVC invoque l’interpréteur « python » présent dans le PATH. La "
         "présence d’un Python système/conda pouvait masquer l’environnement "
         "virtuel du projet. La cible « make dvc-repro » place explicitement "
         "l’environnement virtuel en tête du PATH pour fiabiliser l’exécution.")

    h2(doc, "9.4 Inférence à observation unique")
    body(doc,
         "Certaines variables (moyennes mobiles, indice température-humidité) "
         "supposent un historique absent lors d’une prédiction unitaire. J’ai "
         "retenu des approximations explicites et documentées ; leur impact est "
         "négligeable, car ces variables ont une importance quasi nulle dans le "
         "modèle.")


def conclusion(doc):
    h1(doc, "10. Conclusion et perspectives")
    body(doc,
         "Ce projet couvre les cinq compétences visées par des livrables "
         "fonctionnels et vérifiables : une API REST sécurisée et documentée (C9), "
         "une application accessible qui l’intègre (C10), un dispositif de "
         "monitoring qui collecte, alerte et restitue (C11), une suite de tests "
         "automatisés avec barrière de qualité (C12) et une chaîne de livraison "
         "continue MLOps (C13).")
    body(doc,
         "Le fil conducteur de mon travail a été la rigueur sur les données : "
         "diagnostiquer la cible avant de modéliser, éviter toute fuite, garder un "
         "pipeline reproductible. C’est ce qui a transformé un jeu de données en "
         "apparence inexploitable en une solution qui fonctionne.")
    body(doc, "Je vois trois prolongements à ce travail :")
    bullet(doc, "Câbler un déploiement cloud effectif (la livraison continue "
                "s’arrête aujourd’hui à la publication de l’image et de l’artefact).")
    bullet(doc, "Proposer une interface strictement conforme au RGAA, réutilisant "
                "la même API.")
    bullet(doc, "Enrichir le monitoring d’un suivi de la performance prédictive "
                "dès que des valeurs réelles d’AQI deviennent disponibles, pour "
                "boucler l’amélioration itérative du modèle.")


def annexes(doc):
    h1(doc, "11. Annexes")

    h2(doc, "Annexe A : Référentiel des routes de l’API")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, txt in enumerate(["Route", "Méthode", "Authentification", "Rôle"]):
        table.rows[0].cells[i].paragraphs[0].add_run(txt).bold = True
    rows = [
        ("/health", "GET", "Non", "État du service"),
        ("/model/info", "GET", "Non", "Métadonnées et métriques"),
        ("/predict", "POST", "Oui", "Prédiction unitaire"),
        ("/predict/batch", "POST", "Oui", "Prédiction par lot"),
        ("/metrics", "GET", "Non", "Métriques Prometheus"),
        ("/docs, /redoc", "GET", "Non", "Documentation OpenAPI"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    doc.add_paragraph()

    h2(doc, "Annexe B : Glossaire technique")
    glossaire = [
        ("AQI", "Air Quality Index, indice normalisé de qualité de l’air."),
        ("API REST", "Interface standardisée permettant à des programmes de "
                     "communiquer via HTTP."),
        ("CI/CD", "Intégration et livraison continues : automatisation des tests "
                  "et de la livraison."),
        ("Data drift", "Dérive des données : évolution de la distribution des "
                       "données par rapport à l’entraînement."),
        ("DVC", "Data Version Control : versionnement des données et des modèles, "
                "et orchestration de pipelines."),
        ("MLOps", "Application des bonnes pratiques d’ingénierie logicielle au "
                  "cycle de vie des modèles."),
        ("R²", "Coefficient de détermination : qualité d’ajustement d’un modèle "
               "de régression."),
        ("RGAA / WCAG", "Référentiels d’accessibilité des contenus numériques."),
    ]
    for terme, deff in glossaire:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{terme} : ").bold = True
        p.add_run(deff)


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    for lvl, size in [("Heading 1", 14), ("Heading 2", 11.5), ("Heading 3", 10.5)]:
        st = doc.styles[lvl]
        st.font.color.rgb = ACCENT
        st.font.size = Pt(size)

    add_footer(doc)
    cover(doc)
    sommaire(doc)
    introduction(doc)
    donnees(doc)
    architecture(doc)
    comp_c9(doc)
    comp_c10(doc)
    comp_c11(doc)
    comp_c12(doc)
    comp_c13(doc)
    difficultes(doc)
    conclusion(doc)
    annexes(doc)
    set_update_fields(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Rapport écrit : {OUT}")


if __name__ == "__main__":
    build()
